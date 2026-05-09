from docx import Document
from docx.shared import Inches
from pandas.core.window import doc


def generate_report(df):
    doc = Document()
    doc.add_heading("Sales Data Analysis Report", level=1)

    # Project Objective
    doc.add_paragraph("Analyzes sales performance using Python and Pandas")

    # Dateset Info
    doc.add_heading("Dataset Infromation", level=2)
    doc.add_paragraph(f'Total Rows:{df.shape[1]}')

   # Business KPIs
    doc.add_heading("Business KPIs", level=2)

    # Total Sales
    total_sales = df['Sales'].sum()
    total_profit = df['Profit'].sum()
    total_revenue = df['Revenue'].sum()

    # add in para
    doc.add_paragraph(f'Total Sales :{total_sales}')

    doc.add_paragraph(f'Total Profit :{total_profit}')

    doc.add_paragraph(f'Total Revenue :{total_revenue}')

    # Insert Chart
    doc.add_heading('Revenue Chart',level=2)

    # add pic
    doc.add_picture('Outputs/Charts/Sales_by_Month.png', width=Inches(5))
    doc.add_picture('Outputs/Charts/Sales_By_Region.png', width=Inches(5))
    doc.add_picture('Outputs/Charts/Sales_by_Year.png', width=Inches(5))

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph('Business sales performance was analyzed successfully')

    doc.save("Outputs/Reports/Sales_by_Month.docx")

    print("Word Report is generated successfully")


